import os
import re
import requests
import pandas as pd
import streamlit as st
from bs4 import BeautifulSoup
import docx
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from io import BytesIO

# ==========================================
# Meta-pakker (KAN / SKAL justeres løbende)
# ==========================================
META_PACKAGES = {
    "Meta Ads": {
        "Bronze": {
            "annual_campaigns": 2,
            "ads_per_campaign": 2,
            "horizon_months": 12,
            "features": {
                "remarketing": True,
                "pixel": True,
                "extended_audiences": True,
                "instagram": True,
                "messenger_audience_network": False,
                "conversion_tracking": True,
                "animated_ads": False,
                "story_ads": False,
                "capi_gateway": False,
                "funnel_flows": False,
                "lead_ads": True,
                "dynamic_product_ads": False,
                "meta_shopping": False,
            },
            "description": "Meta Ads Bronze – basis Meta-setup med få kampagner og uden avancerede formater.",
        },
        "Sølv": {
            "annual_campaigns": 3,
            "ads_per_campaign": 2,
            "horizon_months": 12,
            "features": {
                "remarketing": True,
                "pixel": True,
                "extended_audiences": True,
                "instagram": True,
                "messenger_audience_network": True,
                "conversion_tracking": True,
                "animated_ads": True,
                "story_ads": True,
                "capi_gateway": False,
                "funnel_flows": False,
                "lead_ads": True,
                "dynamic_product_ads": False,
                "meta_shopping": False,
            },
            "description": "Meta Ads Sølv – flere kampagner og simple funnel-elementer.",
        },
        "Guld": {
            "annual_campaigns": 4,
            "ads_per_campaign": 3,
            "horizon_months": 12,
            "features": {
                "remarketing": True,
                "pixel": True,
                "extended_audiences": True,
                "instagram": True,
                "messenger_audience_network": True,
                "conversion_tracking": True,
                "animated_ads": True,
                "story_ads": True,
                "capi_gateway": True,
                "funnel_flows": True,
                "lead_ads": True,
                "dynamic_product_ads": False,
                "meta_shopping": False,
            },
            "description": "Meta Ads Guld – stærkere funnel med flere kampagner og tracking.",
        },
        "Platin": {
            "annual_campaigns": 6,
            "ads_per_campaign": 4,
            "horizon_months": 12,
            "features": {
                "remarketing": True,
                "pixel": True,
                "extended_audiences": True,
                "instagram": True,
                "messenger_audience_network": True,
                "conversion_tracking": True,
                "animated_ads": True,
                "story_ads": True,
                "capi_gateway": True,
                "funnel_flows": True,
                "lead_ads": True,
                "dynamic_product_ads": False,
                "meta_shopping": False,
            },
            "description": "Meta Ads Platin – større setup med flere kampagner og fuld funnel.",
        },
        "Platin+": {
            "annual_campaigns": 10,
            "ads_per_campaign": 5,
            "horizon_months": 12,
            "features": {
                "remarketing": True,
                "pixel": True,
                "extended_audiences": True,
                "instagram": True,
                "messenger_audience_network": True,
                "conversion_tracking": True,
                "animated_ads": True,
                "story_ads": True,
                "capi_gateway": True,
                "funnel_flows": True,
                "lead_ads": True,
                "dynamic_product_ads": False,
                "meta_shopping": False,
            },
            "description": "Meta Ads Platin+ – max antal kampagner og avanceret funnel.",
        },
    },
    "Meta Ads Webshop": {
        "Bronze": {
            "annual_campaigns": 6,
            "ads_per_campaign": 3,
            "horizon_months": 12,
            "features": {
                "remarketing": True,
                "pixel": True,
                "extended_audiences": True,
                "instagram": True,
                "messenger_audience_network": False,
                "conversion_tracking": True,
                "animated_ads": True,  # begrænset antal
                "story_ads": True,
                "capi_gateway": False,
                "funnel_flows": False,
                "lead_ads": False,
                "dynamic_product_ads": True,
                "meta_shopping": False,
            },
            "description": "Meta Ads Webshop Bronze – basis webshop-setup med DPA og få kampagner.",
        },
        "Sølv": {
            "annual_campaigns": 8,
            "ads_per_campaign": 4,
            "horizon_months": 12,
            "features": {
                "remarketing": True,
                "pixel": True,
                "extended_audiences": True,
                "instagram": True,
                "messenger_audience_network": True,
                "conversion_tracking": True,
                "animated_ads": True,
                "story_ads": True,
                "capi_gateway": False,
                "funnel_flows": False,
                "lead_ads": False,
                "dynamic_product_ads": True,
                "meta_shopping": False,
            },
            "description": "Meta Ads Webshop Sølv – flere kampagner og flere animerede ads.",
        },
        "Guld": {
            "annual_campaigns": 12,
            "ads_per_campaign": 5,
            "horizon_months": 12,
            "features": {
                "remarketing": True,
                "pixel": True,
                "extended_audiences": True,
                "instagram": True,
                "messenger_audience_network": True,
                "conversion_tracking": True,
                "animated_ads": True,
                "story_ads": True,
                "capi_gateway": True,
                "funnel_flows": True,
                "lead_ads": False,
                "dynamic_product_ads": True,
                "meta_shopping": True,
            },
            "description": "Meta Ads Webshop Guld – fuld webshop-funnel inkl. Meta Shopping.",
        },
        "Platin": {
            "annual_campaigns": 16,
            "ads_per_campaign": 5,
            "horizon_months": 12,
            "features": {
                "remarketing": True,
                "pixel": True,
                "extended_audiences": True,
                "instagram": True,
                "messenger_audience_network": True,
                "conversion_tracking": True,
                "animated_ads": True,
                "story_ads": True,
                "capi_gateway": True,
                "funnel_flows": True,
                "lead_ads": False,
                "dynamic_product_ads": True,
                "meta_shopping": True,
            },
            "description": "Meta Ads Webshop Platin – højt tryk på kampagner og DPA.",
        },
        "Platin+": {
            "annual_campaigns": 20,
            "ads_per_campaign": 6,
            "horizon_months": 12,
            "features": {
                "remarketing": True,
                "pixel": True,
                "extended_audiences": True,
                "instagram": True,
                "messenger_audience_network": True,
                "conversion_tracking": True,
                "animated_ads": True,
                "story_ads": True,
                "capi_gateway": True,
                "funnel_flows": True,
                "lead_ads": False,
                "dynamic_product_ads": True,
                "meta_shopping": True,
            },
            "description": "Meta Ads Webshop Platin+ – maks setup med mange kampagner og DPA.",
        },
    },
    "Meta Ads Performance": {
        "Performance": {
            "annual_campaigns": 4,
            "ads_per_campaign": 3,
            "horizon_months": 4,  # korttidsaftale
            "features": {
                "remarketing": True,
                "pixel": True,
                "extended_audiences": True,
                "instagram": True,
                "messenger_audience_network": False,
                "conversion_tracking": True,
                "animated_ads": True,
                "story_ads": True,
                "capi_gateway": False,
                "funnel_flows": True,
                "lead_ads": True,
                "dynamic_product_ads": False,
                "meta_shopping": False,
            },
            "description": "Meta Ads Performance – 4 måneders performancefokus.",
        },
        "Performance Plus": {
            "annual_campaigns": 6,
            "ads_per_campaign": 3,
            "horizon_months": 4,
            "features": {
                "remarketing": True,
                "pixel": True,
                "extended_audiences": True,
                "instagram": True,
                "messenger_audience_network": True,
                "conversion_tracking": True,
                "animated_ads": True,
                "story_ads": True,
                "capi_gateway": False,
                "funnel_flows": True,
                "lead_ads": True,
                "dynamic_product_ads": True,
                "meta_shopping": False,
            },
            "description": "Meta Ads Performance Plus – mere tryk og flere kampagner på 4 måneder.",
        },
        "Performance Max": {
            "annual_campaigns": 8,
            "ads_per_campaign": 3,
            "horizon_months": 4,
            "features": {
                "remarketing": True,
                "pixel": True,
                "extended_audiences": True,
                "instagram": True,
                "messenger_audience_network": True,
                "conversion_tracking": True,
                "animated_ads": True,
                "story_ads": True,
                "capi_gateway": True,
                "funnel_flows": True,
                "lead_ads": True,
                "dynamic_product_ads": True,
                "meta_shopping": False,
            },
            "description": "Meta Ads Performance Max – maksimalt setup i en kort periode.",
        },
    },
}

#
# =============================
# App config
# =============================
st.set_page_config(page_title="Meta Ads Strategi", layout="wide")
st.title("Meta Ads Strategi")


#
# =============================
# Sidebar (inputs)
# =============================
with st.sidebar:
    st.header("Indstillinger")
    api_key = st.secrets.get("OPENAI_API_KEY", None)
    model = st.selectbox(
        "Model",
        ["gpt-5", "gpt-4o", "gpt-4o-mini"],
        index=0
    )
    st.divider()

    # Pakke-vælger i stedet for frit antal kampagner
    package_family = st.selectbox("Meta produktpakke", list(META_PACKAGES.keys()), index=0)
    tier_options = list(META_PACKAGES[package_family].keys())
    package_tier = st.selectbox("Pakke-niveau", tier_options, index=0)

    current_pkg = META_PACKAGES[package_family][package_tier]
    base_campaigns = current_pkg["annual_campaigns"]
    ads_per_campaign = current_pkg["ads_per_campaign"]
    horizon_months = current_pkg["horizon_months"]
    package_features = current_pkg["features"]
    package_description = current_pkg["description"]

    extra_campaigns = st.number_input(
        "Ekstra kampagner udover pakken",
        min_value=0,
        max_value=20,
        value=0,
        step=1
    )
    total_campaigns = base_campaigns + extra_campaigns

    st.caption(
        f"Pakke: {package_family} {package_tier} – {base_campaigns} kampagner årligt, "
        f"{ads_per_campaign} annoncer pr. kampagne, horisont: {horizon_months} mdr."
    )

    monthly_budget = st.number_input("Månedligt budget (DKK)", min_value=0, step=1000, value=0)

#
# =============================
# Main inputs
# =============================
col1, col2 = st.columns(2)
with col1:
    customer_name = st.text_input("Kundenavn")
    website = st.text_input("Website (https://…)")
    important_subpages_raw = st.text_area("Vigtige undersider (fulde URLs, én per linje)", height=120)

    # ==========================================
    # Background scraping
    # ==========================================
    if website:
        # --- Helper: simple_scrape ---
        def simple_scrape(url: str, timeout_sec: int = 10) -> dict:
            """Let scraping: forsiden + op til 3 interne links (titel, meta, H1/H2, text)."""
            out = {"homepage": {}, "samples": []}
            if not url or not url.startswith("http"):
                return out
            try:
                r = requests.get(url, timeout=timeout_sec, headers={"User-Agent": "Mozilla/5.0"})
                soup = BeautifulSoup(r.text, "lxml")
                title = soup.title.text.strip() if soup.title else ""
                md = soup.find("meta", {"name": "description"})
                meta = md.get("content", "").strip() if md else ""
                h1 = [h.get_text(" ", strip=True) for h in soup.find_all("h1")][:5]
                h2 = [h.get_text(" ", strip=True) for h in soup.find_all("h2")][:8]
                out["homepage"] = {"title": title, "meta": meta, "h1": h1, "h2": h2}

                links = []
                for a in soup.find_all("a", href=True):
                    href = a["href"].strip()
                    if href.startswith("http") and url.split("/")[2] in href:
                        links.append(href)
                    elif href.startswith("/"):
                        base = url.rstrip("/")
                        links.append(base + href)
                    if len(links) >= 3:
                        break

                for link in links:
                    try:
                        r2 = requests.get(link, timeout=timeout_sec, headers={"User-Agent": "Mozilla/5.0"})
                        s2 = BeautifulSoup(r2.text, "lxml")
                        title2 = s2.title.text.strip() if s2.title else ""
                        md2 = s2.find("meta", {"name": "description"})
                        meta2 = md2.get("content", "").strip() if md2 else ""
                        h1_2 = [h.get_text(" ", strip=True) for h in s2.find_all("h1")][:3]
                        paragraphs = [p.get_text(" ", strip=True) for p in s2.find_all("p")]
                        lis = [li.get_text(" ", strip=True) for li in s2.find_all("li")]
                        combined = paragraphs + lis
                        combined = [c for c in combined if len(c) > 40]
                        text2 = "\n".join(combined[:20])

                        out["samples"].append({
                            "url": link,
                            "title": title2,
                            "meta": meta2,
                            "h1": h1_2,
                            "text": text2
                        })
                    except Exception:
                        continue
            except Exception:
                pass
            return out

        prev = st.session_state.get("scraped_site", {})
        if prev.get("_url") != website:
            st.session_state["scraped_site"] = {
                "_url": website,
                "data": simple_scrape(website)
            }

    # --- Scrape a single page (full URL) for title/meta/h1/text ---
    def scrape_single_page(url: str, timeout_sec: int = 10) -> dict:
        out = {"url": url, "title": "", "meta": "", "h1": [], "text": ""}
        if not url or not url.startswith("http"):
            return out
        try:
            r = requests.get(url, timeout=timeout_sec, headers={"User-Agent": "Mozilla/5.0"})
            s = BeautifulSoup(r.text, "lxml")

            out["title"] = s.title.text.strip() if s.title else ""
            md = s.find("meta", {"name": "description"})
            out["meta"] = md.get("content", "").strip() if md else ""
            out["h1"] = [h.get_text(" ", strip=True) for h in s.find_all("h1")][:5]

            # NEW: extract visible text (A-level scraping)
            paragraphs = [p.get_text(" ", strip=True) for p in s.find_all("p")]
            lis = [li.get_text(" ", strip=True) for li in s.find_all("li")]

            combined = paragraphs + lis
            combined = [c for c in combined if len(c) > 40]  # filter noise
            out["text"] = "\n".join(combined[:20])  # cap to 20 elements
        except Exception:
            pass
        return out

    if important_subpages_raw:
        sub_urls = [u.strip() for u in important_subpages_raw.splitlines() if u.strip()]
        prev = st.session_state.get("scraped_subpages", {})
        if prev.get("_urls") != sub_urls:
            st.session_state["scraped_subpages"] = {
                "_urls": sub_urls,
                "data": [scrape_single_page(u) for u in sub_urls]
            }

    other_info = st.text_area("Egne idéer / Anden vigtig info", height=120)
    competitors_raw = st.text_area("Konkurrenter (én per linje)", height=120)

with col2:
    xpect_doc = st.file_uploader("Xpect (DOCX/TXT/PDF)", type=["docx", "txt", "pdf"])
    ad_data_file = st.file_uploader("Eksisterende data fra annoncekonto (CSV/Excel)", type=["csv", "xlsx"])

generate_btn = st.button("Generer strategi")

#
# =============================
# Helpers
# =============================

def safe_read_xpect(file, manual_text=""):
    """Returnér ren tekst fra Xpect upload (eller tom)."""
    if file is not None:
        ext = os.path.splitext(file.name.lower())[1]
        try:
            if ext == ".txt":
                return file.read().decode("utf-8", errors="ignore")
            if ext == ".pdf":
                data = file.read()
                txt = data.decode("latin-1", errors="ignore")
                return re.sub(r"[^A-Za-zÆØÅæøå0-9 ,.\-–:;()\n]+", " ", txt)
            if ext == ".docx":
                doc = docx.Document(file)
                return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        except Exception:
            return manual_text.strip()
    return manual_text.strip()


def summarize_ad_account(df: pd.DataFrame) -> str:
    """Kort tekstopsummering af eksisterende kontodata (hvis uploadet)."""
    try:
        cols = [c.lower() for c in df.columns]
        out = []

        def col(name): return df.columns[cols.index(name)]

        if "spend" in cols or "cost" in cols:
            spend_col = col("spend") if "spend" in cols else col("cost")
            total_spend = df[spend_col].sum(numeric_only=True)
            out.append(f"Samlet spend: {total_spend:,.0f} DKK")
        if "impressions" in cols:
            out.append(f"Impressions: {df[col('impressions')].sum(numeric_only=True):,.0f}")
        if "clicks" in cols:
            out.append(f"Klik: {df[col('clicks')].sum(numeric_only=True):,.0f}")
        if "conversions" in cols:
            out.append(f"Konverteringer: {df[col('conversions')].sum(numeric_only=True):,.0f}")
        if "cpa" in cols:
            out.append(f"Gns. CPA: {df[col('cpa')].mean(numeric_only=True):,.2f} DKK")
        return " | ".join(out) if out else "Ingen standard KPI-kolonner genkendt – data vedlagt som rå bilag."
    except Exception:
        return "Kunne ikke opsummere kontodata – behandles som rå bilag."



def sanitize(txt: str) -> str:
    """Fjern markdown, dobbelte mellemrum og bevar bullets med '•'."""
    if not txt:
        return ""
    txt = txt.replace("**", "")
    txt = re.sub(r"^#{1,6}\s*", "", txt, flags=re.MULTILINE)
    txt = txt.replace("* ", "• ")
    txt = txt.replace("\t", " ")
    txt = re.sub(r" {2,}", " ", txt)
    return txt.strip()

# --- Helper: Format overskrifter som H2 + fede underoverskrifter ---
def format_headings(text: str) -> str:
    """
    Finder strategiens kendte overskrifter og gør dem til rigtige H2-overskrifter i Streamlit.
    Gør samtidig centrale underoverskrifter (Forretningen:, Produkter/ services: osv.)
    tydelige med fed skrift.
    """
    if not text:
        return ""

    # Hovedoverskrifter (sektionstitler)
    headings = [
        "Strategi",
        "Agenda",
        "Forretningen",
        "Introduktion til Meta strategi",
        "Målsætninger & KPI’er",
        "Målgruppe",
        "Kampagneplan & budgetplan",
        "USP’er & budskaber",
        "Content",
        "Next",
    ]
    for h in headings:
        # Erstat linjer der KUN består af overskriften med en H2 + fed
        text = re.sub(
            rf"^{re.escape(h)}$",
            f"<h2><b>{h}</b></h2>",
            text,
            flags=re.MULTILINE,
        )

    # Underoverskrifter, der skal stå som fede linjer men ikke som H2
    subheadings = [
        "Forretningen:",
        "Produkter/ services:",
        "Anbefalet content:",
        "Spørgsmål:",
        "Produkt/service USP’er:",
        "Brand-USP’er:",
        "3-8 spørgsmål der sikrer en mere præcis forståelse:",
    ]
    for sh in subheadings:
        text = re.sub(
            rf"^{re.escape(sh)}$",
            f"<b>{sh}</b>",
            text,
            flags=re.MULTILINE,
        )

    return text


def build_campaign_table(text: str) -> pd.DataFrame:
    """
    Uddrag kampagneoplysninger fra eksekveringsoutputtet til oversigtstabel.
    Matcher formatet:

    Kampagne X: [navn]
    Formål: ...
    Indhold: ...
    Målgruppe: ...
    Budget: ...
    Periode: ...
    """
    campaigns = []
    current = {}
    for line in text.splitlines():
        line = line.strip()
        if not line:
            continue

        lower = line.lower()
        if re.match(r"^kampagne\s+\d+", lower):
            # Ny kampagneblok (kun linjer som "Kampagne 1", "Kampagne 2" osv. – IKKE "Kampagneoversigt")
            if current:
                campaigns.append(current)
            current = {"Kampagne": line}
        elif lower.startswith("formål:"):
            current["Formål"] = line.split(":", 1)[1].strip()
        elif lower.startswith("målgruppe:"):
            current["Målgruppe"] = line.split(":", 1)[1].strip()
        elif lower.startswith("budget:"):
            current["Budget"] = line.split(":", 1)[1].strip()
        elif lower.startswith("periode:"):
            current["Periode"] = line.split(":", 1)[1].strip()

    if current:
        campaigns.append(current)

    rows = []
    for c in campaigns:
        rows.append({
            "Kampagne": c.get("Kampagne", ""),
            "Formål": c.get("Formål", ""),
            "Målgruppe": c.get("Målgruppe", ""),
            "Periode": c.get("Periode", ""),
            "Budget": c.get("Budget", ""),
        })

    return pd.DataFrame(rows)


# =============================
# Budgetplan-table builder
# =============================
def build_budget_table(text: str) -> pd.DataFrame:
    """
    Uddrag budgetplan pr. måned fra eksekveringsoutputtet til oversigtstabel.
    Matcher formatet:

    Budgetplan:
    Januar
    Budget (DKK): ca. 1.200
    Fokus: ...
    Begrundelse: ...
    """
    lines = [ln.strip() for ln in text.splitlines()]
    budget_idx = None
    months = {
        "Januar",
        "Februar",
        "Marts",
        "April",
        "Maj",
        "Juni",
        "Juli",
        "August",
        "September",
        "Oktober",
        "November",
        "December",
    }
    for i, ln in enumerate(lines):
        if ln.lower().startswith("budgetplan"):
            budget_idx = i
            break
    if budget_idx is None:
        return pd.DataFrame()

    rows = []
    i = budget_idx + 1
    n = len(lines)

    while i < n:
        # Skip tomme linjer
        while i < n and not lines[i]:
            i += 1
        if i >= n:
            break

        month = lines[i]
        # Stop budgetsektionen, hvis linjen ikke er en måned
        if month not in months:
            break
        budget = ""
        fokus = ""
        begrundelse = ""

        if i + 1 < n and lines[i + 1].lower().startswith("budget"):
            budget = lines[i + 1].split(":", 1)[1].strip()
        if i + 2 < n and lines[i + 2].lower().startswith("fokus"):
            fokus = lines[i + 2].split(":", 1)[1].strip()
        if i + 3 < n and lines[i + 3].lower().startswith("begrundelse"):
            begrundelse = lines[i + 3].split(":", 1)[1].strip()

        rows.append(
            {
                "Måned": month,
                "Budget (DKK)": budget,
                "Fokus": fokus,
                "Begrundelse": begrundelse,
            }
        )

        # Hop frem til næste blok (vi antager tom linje mellem måneder)
        i += 4
        while i < n and lines[i]:
            i += 1

    return pd.DataFrame(rows)


#
# =============================
# OpenAI wrapper (med streaming)
# =============================

def run_gpt(prompt: str, api_key: str, model: str, max_tokens: int = 2000, stream_placeholder=None) -> str:
    """
    OpenAI wrapper for konsistente completion-kald, med mulighed for streaming til Streamlit.
    """
    from openai import OpenAI
    client = OpenAI(api_key=api_key)
    try:
        model_map = {
            "gpt-5": "gpt-5",
            "gpt-4o": "gpt-4o",
            "gpt-4o-mini": "gpt-4o-mini"
        }
        temp_map = {
            "gpt-5": 1.0,
            "gpt-4o": 0.0,
            "gpt-4o-mini": 0.0
        }
        chosen_model = model_map.get(model, model)
        temperature_value = temp_map.get(model, 0.0)

        system_msg = "Du er en senior Meta Ads strategist. Svar altid i ren tekst uden markdown."

        # Ingen streaming: klassisk kald
        if stream_placeholder is None:
            response = client.chat.completions.create(
                model=chosen_model,
                messages=[
                    {"role": "system", "content": system_msg},
                    {"role": "user", "content": prompt}
                ],
                temperature=temperature_value,
            )
            content = response.choices[0].message.content.strip() if response.choices else ""
            if not content:
                return "[AI-ERROR] Modellen returnerede tomt svar."
            return content

        # Streaming
        full_text = ""
        stream = client.chat.completions.create(
            model=chosen_model,
            messages=[
                {"role": "system", "content": system_msg},
                {"role": "user", "content": prompt}
            ],
            temperature=temperature_value,
            stream=True,
        )
        for chunk in stream:
            if not chunk.choices:
                continue
            delta = chunk.choices[0].delta.content or ""
            if not delta:
                continue
            full_text += delta
            # Brug Markdown linebreaks for pæn visning mens der streams
            stream_placeholder.markdown(full_text.replace("\n", "  \n"))

        full_text = full_text.strip()
        if not full_text:
            return "[AI-ERROR] Modellen returnerede tomt streamsvar."
        return full_text
    except Exception as e:
        return f"[AI-ERROR] {e}"


#
# =============================
# Prompts (2-kalds flow)
# =============================

def build_context():
    # Xpect
    xpect_text = safe_read_xpect(xpect_doc, "")

    # Eksisterende kontodata
    ad_account_summary = ""
    if ad_data_file is not None:
        try:
            if ad_data_file.name.endswith(".csv"):
                df = pd.read_csv(ad_data_file)
            else:
                df = pd.read_excel(ad_data_file)
            ad_account_summary = summarize_ad_account(df)
        except Exception:
            ad_account_summary = "Kunne ikke læse filen – ignoreres i analysen."

    # Let website-scrape (valgfrit input til prompten)
    site = st.session_state.get("scraped_site", {}).get("data", {})

    # Vigtige undersider (fra brugerinput)
    user_subpages = st.session_state.get("scraped_subpages", {}).get("data", [])

    # Kampagner (samlet antal)
    total = max(1, int(total_campaigns))

    # Parse competitors
    competitors = [c.strip() for c in (competitors_raw or "").splitlines() if c.strip()]

    return {
        "customer_name": customer_name,
        "website": website,
        "monthly_budget": monthly_budget,
        "other_info": other_info,
        "xpect": xpect_text,
        "ad_account_summary": ad_account_summary,
        "site": site,
        "total_campaigns": total,
        "competitors": competitors,
        "user_subpages": user_subpages,
        "package_family": package_family,
        "package_tier": package_tier,
        "base_campaigns": base_campaigns,
        "extra_campaigns": extra_campaigns,
        "ads_per_campaign": ads_per_campaign,
        "horizon_months": horizon_months,
        "package_features": package_features,
        "package_description": package_description,
    }


def prompt_strategy_core(ctx: dict) -> str:
    # Brug site/Xpect/undersider som kontekst
    home = ctx["site"].get("homepage", {})
    subpages = ctx.get("user_subpages", []) or []

    # Byg kort tekstblok ud fra vigtige undersider
    subpage_summaries = []
    for s in subpages:
        if not s:
            continue
        url = s.get("url", "")
        title = s.get("title", "")
        txt = (s.get("text", "") or "").replace("\n", " ")
        if len(txt) > 260:
            txt = txt[:260] + "..."
        label = title or url
        if not label:
            continue
        subpage_summaries.append(f"- {label}: {txt}")
    subpages_block = "\n".join(subpage_summaries[:5])

    sub_titles = [s.get("title", "") for s in ctx.get("user_subpages", []) if s.get("title")]
    sample_titles = ", ".join(
        ([s.get("title", "") for s in ctx["site"].get("samples", []) if s.get("title")] + sub_titles)[:6]
    )

    # Hårdt prioriterede inputfelter
    other_info = ctx.get("other_info") or "(tom)"
    xpect_excerpt = (ctx.get("xpect") or "")[:800]
    ad_account_summary = ctx.get("ad_account_summary") or "Ingen data"

    return f"""
Du er senior Meta Ads-strateg og skal levere den INDLEDENDE STRATEGI til et kundemøde på DANSK som ren tekst med bullets (•) hvor det er naturligt.

OUTPUTTET SKAL FØLGE DENNE SKABELON 1:1 I STRUKTUR, RÆKKEFØLGE OG OVERSKRIFTER (ingen ekstra sektioner, ingen ændrede overskrifter):

Agenda
Forretningen
Introduktion til Meta strategi
Målsætninger & KPI’er
Målgruppe

Du SKAL aktivt bruge og PRIORITERE følgende input i strategien:

1) Egne idéer / Anden vigtig info (HØJESTE PRIORITET)
   • Alt i strategien (forretning, strategi, målsætninger, målgrupper) skal tydeligt afspejle dette felt.
   • Rå tekst:
   {other_info}

2) Vigtige undersider (fulde URLs)
   • De vigtige undersider viser hvad kunden selv fremhæver ift. produkter/ydelser og fokusområder.
   • Brug dem til at forstå sortiment, services, argumenter og eventuelle sæsoner.
   • Kort udtræk fra undersiderne:
{subpages_block or '- (ingen undersider angivet)'}

3) Xpect (DOCX/TXT/PDF)
   • Xpect er strategibriefet og skal respekteres.
   • Brug uddraget som ekstra kontekst til tone, fokus, mål og constraints.
   • Xpect (uddrag, maks. 800 tegn):
{xpect_excerpt}

4) Eksisterende data fra annoncekonto (CSV/Excel)
   • Kontodata fortæller noget om historisk spend, reach og eventuelle konverteringer.
   • Brug dette til at nuancere formuleringer om målsætninger (fx “vi starter med at bygge baseline”, “der er allerede dokumenteret performance” osv.).
   • Kort opsummering:
   {ad_account_summary}

Derudover får du overblik over website og øvrige undersider:
    Kunde: {ctx.get('customer_name') or '(ukendt)'}
    Website: {ctx.get('website') or '(ukendt)'}
    Månedligt budget: {ctx.get('monthly_budget')} DKK
    Website (titel/meta/h1 fra forside): {home.get('title', '')} | {home.get('meta', '')} | {', '.join(home.get('h1', [])[:3])}
    Eksempler på øvrige undersider (titler): {sample_titles}

VIGTIGT:
• Feltet “Egne idéer / Anden vigtig info” er HØJESTE prioritet og SKAL tydeligt kunne genkendes i alle frie tekstafsnit.
• “Vigtige undersider” og Xpect skal bruges aktivt til at forstå produkter, services, sæson og positionering – de må IKKE ignoreres.
• Eksisterende data fra annoncekonto skal bruges til at beskrive niveauet af historik og påvirke hvor forsigtig/ambitiøs du er i formuleringerne om målsætninger og læringsfaser.
• Du må IKKE ændre på skabelonens overskrifter eller rækkefølge.
• Du må kun tilføje korte forklarende sætninger – selve FORSLAG-teksten og spørgsmålene er KUN til inspiration til dig og må ikke fremgå af det endelige resultat.

STRUKTUR OG INDHOLD (PRÆCIS SOM SKABELONEN):

1) Agenda
    • Linje 1: præcis overskrift: Agenda
    • Derefter følgende punkter, én pr. linje med bullet "•" (ingen ekstra, ingen færre):
      • Velkomst & Introduktion
      • Jeres forretning
      • Introduktion til Meta strategi
      • Målsætninger & KPI’er
      • Målgruppe
      • Kampagne- og budgetplan
      • USP’er & budskaber
      • Content
      • Next

2) Forretningen
    • Overskriften skal være alene på en linje: Forretningen
    • Skriv 2–5 linjer, der kort opsummerer forretningen baseret på konteksten (hvad de sælger, hvem de sælger til, geografi, B2B/B2C, salgsvej, udfordringer).
    • Her skal du tydeligt afspejle specialistens input fra “Egne idéer / Anden vigtig info” og indsigter fra både Xpect og vigtige undersider.

    FORSLAG til inspiration:
    Forretningen:
    Beskriv jeres virksomhed/ hvem I er
    Hvad er jeres vigtigste styrker – og svagheder?
    Hvad tjener I pr. kunde (livstidsværdi)?
    Hvordan ser en typisk kunderejse ud? Hvordan kontakter kunderne jer? 
    Hvordan varierer jeres salg hen over året (sæson)?

    Produkter/ services:
    Hvilke produkter/services er vigtigst, vi fokuserer på?
    Hvilke produkter/services giver jer størst indtjening?
    Hvad driver typisk et køb – pris, kvalitet, brand eller noget andet?
    Har I specifikke kollektioner, kampagner eller nyheder, vi skal planlægge efter i løbet af året?

3) Introduktion til Meta strategi
    • Overskriften skal være alene på en linje: Introduktion til Meta strategi
    • Skriv 3–6 linjer, der forklarer strategiens overordnede tilgang – baseret på eksemplet i skabelonen:
        Annonceringen skal opbygges med to spor:
        Det ene område fokuserer på branding og at opbygge en kvalificeret målgruppe.
        Det andet har fokus på leadgenerering/køb, hvor annoncer leder ind i det relevante flow (fx nyhedsbrevsflow, leadformularer, webshop).
        Formålet er at skabe synergi mellem Meta og øvrige kanaler og modne målgruppen frem mod konvertering.
    • Tilpas formuleringerne til kundens virkelighed ved aktivt at bruge:
        – “Egne idéer / Anden vigtig info”
        – Indhold og fokus fra vigtige undersider
        – Xpect-uddraget
        – Kontodata (hvis de indikerer lav/høj spend eller manglende konverteringshistorik)

4) Målsætninger & KPI’er
    • Overskriften skal være alene på en linje: Målsætninger & KPI’er
    • Start med 2–4 linjer, hvor du kort beskriver den overordnede målsætning ud fra konteksten (flere leads, mere online salg, øget brandkendskab osv.).
    • Du SKAL her bruge:
        – Specialistens “Egne idéer / Anden vigtig info”
        – Kontodata-opsummeringen
        – Eventuelle KPI’er eller mål fra Xpect
      til at gøre målsætningerne realistiske og forankrede i virkeligheden.

    FORSLAG til inspiration:
    E-commerce: ROAS, omsætning, CPA, antal køb, pris pr. køb
    Leads/Service: CPL, antal kvalificerede leads, konverteringsrate
    B2B: CPL, bookede møder, video consumption, hjemmesidetrafik
    Awareness: Reach, videovisninger, CTR, besøg på website/sales pages, nye følgere

5) Målgruppe
    • Overskriften skal være alene på en linje: Målgruppe
    • Skriv 2–4 linjer, der kort opsummerer målgruppen baseret på:
        – Vigtige undersider
        – Xpect
        – “Egne idéer / Anden vigtig info”
        – Website-kontekst (forside + evt. samples)
    • Derefter SKAL følgende del stå ordret:

    FORSLAG til inspiration:
    3-8 spørgsmål der sikrer en mere præcis forståelse:
    Hvem er jeres drømmekunde?
    Hvilke motiver driver deres køb?
    Hvorfor vælger de jer fremfor konkurrenterne?
    Hvilke barrierer eller objections kan der være?
    Typiske spørgsmål potentielle kunder stiller?


GENERELLE OUTPUTKRAV:
    • Svar altid i ren tekst uden markdown (ingen #, ingen **).
    • Brug præcis disse overskrifter: “Agenda”, “Forretningen”, “Introduktion til Meta strategi”, “Målsætninger & KPI’er”, “Målgruppe” – hver på sin egen linje.
    • Skriv ALDRIG ordet "FORSLAG" eller "FORSALG" i svaret.
    • Sørg for tydelige linjeskift mellem sektioner og underblokke, så teksten er let at læse højt på et møde.
    • Udnyt konteksten (Egne idéer, vigtige undersider, Xpect, kontodata, website) til de frie tekstlinjer, men opfind ikke konkrete tal eller fakta, kunden ikke har givet.

Før du skriver outputtet, laver du en intern kæde af tanker for at sikre, at de korte, frie tekstafsnit understøtter skabelonens forslag uden at ændre på dem. Start først på selve teksten, når du har planen klar.
"""


def prompt_execution(ctx: dict, strategy_core_text: str) -> str:
    total_campaigns = ctx.get("total_campaigns", 4)

    # Hårdt prioriterede inputfelter til eksekveringsdelen
    other_info = ctx.get("other_info") or "(tom)"
    xpect_excerpt = (ctx.get("xpect") or "")[:800]
    ad_account_summary = ctx.get("ad_account_summary") or "Ingen data"
    subpages = ctx.get("user_subpages", []) or []

    subpage_summaries = []
    for s in subpages:
        if not s:
            continue
        url = s.get("url", "")
        title = s.get("title", "")
        txt = (s.get("text", "") or "").replace("\n", " ")
        if len(txt) > 260:
            txt = txt[:260] + "..."
        label = title or url
        if not label:
            continue
        subpage_summaries.append(f"- {label}: {txt}")
    subpages_block = "\n".join(subpage_summaries[:5])

    prompt = f"""
Du skal nu bygge resten af METASTRATEGI-DOKUMENTET baseret på strategiens kerne (nedenfor).
Svar på DANSK, i ren tekst (ingen markdown, ingen **). Brug bullets "•" kun hvor det er naturligt.

Du SKAL aktivt bruge og prioritere følgende input i kampagneplan, budgetplan, USP’er, content og next steps:

1) Egne idéer / Anden vigtig info (HØJESTE PRIORITET)
   • Kampagner, budget, budskaber og forslag til inspiration SKAL kunne genkendes i dette input.
   • Rå tekst:
   {other_info}

2) Vigtige undersider (fulde URLs)
   • Brug disse sider til at forstå hvilke produkter/ydelser, temaer og sæsoner der skal prioriteres.
   • Kort udtræk:
{subpages_block or '- (ingen undersider angivet)'}

3) Xpect (DOCX/TXT/PDF)
   • Brug Xpect-uddraget til at sikre, at kampagnestruktur, budskaber og prioriteringer hænger sammen med den overordnede brief.
   • Uddrag (maks. 800 tegn):
{xpect_excerpt}

4) Eksisterende data fra annoncekonto (CSV/Excel)
   • Brug kontodata til at nuancere budget og kampagnevalg (fx hvis der mangler historik, hvis spend har været lav/høj, eller hvis konverteringer er ustabile).
   • Kort opsummering:
   {ad_account_summary}

Du skal levere sektioner i PRÆCIS denne rækkefølge med følgende overskrifter, hver på sin egen linje:

Kampagneplan & budgetplan
USP’er & budskaber
Content
Next

VIGTIGT OM KAMPAGNER:
    • Der skal i alt foreslås {total_campaigns} kampagner.
    • Mindst én af kampagnerne SKAL være en Always-On kampagne (tydeligt angivet som Always-On i formål eller periode, fx "Hele året (Always-On)").
    • Kampagnerne skal hænge logisk sammen med målsætning, målgruppe og strategi fra strategiens kerne.
    • “Egne idéer / Anden vigtig info” er højest prioriteret og SKAL kunne genkendes i kampagnernes formål, indhold og målgrupper.
    • Hvis vigtige undersider afslører særlige produkter/temaer, skal disse prioriteres i kampagner og content.
    • Hvis kontodata viser begrænset eller ustabil historik, skal du afspejle dette i, hvordan du beskriver test, læring og risikoniveau i kampagneplanen og budgetplanen.
    • Leadpage SKAL bruges som primær landingsside for B2B-kampagner, ikke webshoppen (medmindre andet fremgår direkte af Egne idéer / Anden vigtig info).
    • Kampagner der retter sig mod forhandlere SKAL indeholde budskaber om margin, levering, ordreflow og hyldeeffekt, når det er relevant i konteksten.

1) Kampagneplan & budgetplan
    • Start med en linje med overskriften: Kampagneplan & budgetplan
    • Tilføj en tom linje og derefter underoverskriften: Kampagneoversigt
    • Derefter beskriver du hver kampagne i følgende format, én kampagne ad gangen:

      Kampagne X: [kort navn på kampagnen]
      Formål: Beskriv kampagnens formål (fx leadgenerering, salg, retargeting, brandopbygning)
      Indhold: Beskriv hvilken type content og budskaber der skal bruges – med tydelig sammenhæng til Egne idéer / Anden vigtig info og til vigtige undersider
      Målgruppe: Beskriv målgruppen for kampagnen
      Budget: Angiv enten cirka-beløb i DKK baseret på månedligt budget eller andel i %
      Periode: Angiv forventet periode (fx Hele året (Always-On), Q1–Q2, kampagnemåneder osv.)

    • Sørg for at:
        – Én kampagne tydeligt er "Always-On"
        – De øvrige kampagner dækker relevante peaks/temaer i løbet af året (sæsoner, udsalg, events osv.), som de fremgår af Egne idéer, Xpect og kontodata
        – Beskrivelserne er korte, konkrete og kan læses højt for kunden.

    • Efter kampagneoversigten tilføjer du en tom linje og skriver: Budgetplan:
    • Under "Budgetplan:" beskriver du årets fordeling måned for måned (Januar til December), inspireret af eksemplet:
        – For hver måned angives:
            Januar
            Budget (DKK): ca. XX (kan være skøn, afledt af månedligt budget)
            Fokus: Beskriv kort hvad budgettet primært bruges til
            Begrundelse: Kort begrundelse for niveauet (fx sæson, opstart, testfase)
        – Brug konkrete ca.-beløb baseret på månedligt budget (fx 2.000 DKK → fordel 1.200 / 400 / 300 / 100).
        – Hvis kontodata viser lav/høj historisk spend, skal du nævne det i begrundelserne (fx “vi starter lavt for at bygge data”, “vi kan skrue op i peaks fordi der er dokumenteret efterspørgsel”).
    • Afslut budgetplanen med 2–3 linjer, der kort opsummerer:
        – hvilke måneder der er særligt tunge
        – hvordan Always-On understøtter resten af året
        – hvordan budgettet kan justeres op/ned, fx baseret på performance og kontodata.

2) USP’er & budskaber
    • Overskriften skal være alene på linjen: USP’er & budskaber
    • Del derefter USP’er i to blokke med præcise underoverskrifter:

      Produkt/service USP’er:
      • Skriv 3–6 bullets i formatet "USP → kundeudbytte", fx:
        Høj kvalitet → Kunden får noget der holder
        Hurtig levering → Mindre ventetid og mindre friktion i købet
        Bredt udvalg → Let at finde det, der passer
      • Brug viden fra vigtige undersider, Xpect og Egne idéer til at formulere USP’erne.

      Brand-USP’er:
      • Skriv 3–6 bullets i samme format, tilpasset konteksten, fx:
        Ekspertise og erfaring → Tryghed i valg af leverandør
        Lokal tilstedeværelse → Let at mødes og få hjælp
        Certificeringer/awards → Dokumenteret faglighed
        Personlig rådgivning → Relation og langvarige samarbejder

    • Fokusér på at koble USP direkte til kundeudbytte og indsigter fra strategiens kerne + de fire prioriterede inputfelter.

3) Content
    • Overskriften skal være alene på linjen: Content
    • Næste linje SKAL være: Anbefalet content:
    • Derefter lister du 6–12 konkrete content-typer i bullets (•), inspireret af eksemplet, fx:
        • Videoer med stærke hooks, der viser problem/løsning
        • Produkt- eller servicebilleder i brug (lifestyle)
        • Testimonials og cases med citater
        • UGC-stil video
        • Casebilleder (før/efter hvis relevant)
        • Grafiske kampagneelementer
      – Sørg for at content-idéerne hænger sammen med:
        • Egne idéer / Anden vigtig info
        • Vigtige undersider (hvilke produkter og temaer der vises)
        • Xpect-briefet
        • De valgte kampagner og deres formål

    • Efter content-listen tilføjer du en linje: Spørgsmål:
    • Under “Spørgsmål:” skriver du 2–4 spørgsmål til kunden om content, inspireret af skabelonen:
        Har I eksisterende content, vi kan bruge?
        Er det muligt at producere anbefalede videoer/ billeder?
      – Disse to spørgsmål SKAL stå ordret.
      – Du må gerne tilføje op til 2 ekstra spørgsmål i samme stil, fx om cases, forhandlerudtalelser eller adgang til materialer, men stadig max 4 spørgsmål i alt.

4) Next
    • Overskriften skal være alene på linjen: Next
    • Næste linje SKAL være: EKSEMPEL:
    • Derefter skriver du 3–6 bullets, der beskriver næste skridt efter mødet, inspireret af skabelonens eksempel:
        Vi færdiggør strategien ud fra mødenoter.
        Vi sender udkast til første kampagne.
        Når alt er godkendt, sættes kampagnerne live.
        Første performanceupdate sendes efter X dage.
      – Du skal her også referere til:
        • Opsætning af tracking/leadflow, så Egne idéer kan realiseres
        • Eventuelle behov for ekstra data/indsigter fra annoncekontoen
        • Eventuel produktion af content, som er nævnt i Content-sektionen.

STRATEGI-KERNEN (KONTEKST):
{strategy_core_text[:4500]}

OUTPUTKRAV:
    • Svar altid i ren tekst uden markdown (ingen #, ingen **).
    • Brug overskrifterne "Kampagneplan & budgetplan", "USP’er & budskaber", "Content" og "Next" præcist som skrevet, hver på sin egen linje.
    • Brug bullets med "•" hvor det er naturligt.
    • Mindst én kampagne SKAL være markeret som Always-On i kampagneplanen.
    • Sørg for tydelige linjeskift og luft mellem underblokke, så teksten er nem at bruge direkte i et Word-dokument og på møder.
    • Udnyt strategiens kerne samt de fire prioriterede inputfelter aktivt – ikke kun som pynt.

Før du skriver selve outputtet, laver du en intern kæde af tanker for at sikre, at kampagnestruktur, budgetplan, USP’er, content og next steps hænger logisk sammen med strategiens kerne OG de prioriterede input (Egne idéer, vigtige undersider, Xpect, kontodata).
"""
    # Fjern evt. rester af gamle "Spørgsmål til kunden"-instrukser hvis de sniger sig ind
    prompt = re.sub(r"Spørgsmål til kunden:?(\s*\+\s*2\s*bullets)?", "", prompt, flags=re.IGNORECASE)
    return prompt


#
# =============================
# DOCX builder (simpel)
# =============================

def add_page_numbers(document: Document):
    for section in document.sections:
        footer = section.footer
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        def _add_field_run(paragraph, field_code):
            r = paragraph.add_run()
            fldChar1 = OxmlElement('w:fldChar')
            fldChar1.set(qn('w:fldCharType'), 'begin')
            r._r.append(fldChar1)
            instr = OxmlElement('w:instrText')
            instr.set(qn('xml:space'), 'preserve')
            instr.text = field_code
            r._r.append(instr)
            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'separate')
            r._r.append(fldChar2)
            r2 = paragraph.add_run()
            fldChar3 = OxmlElement('w:fldChar')
            fldChar3.set(qn('w:fldCharType'), 'end')
            r2._r.append(fldChar3)

        _add_field_run(p, 'PAGE')
        p.add_run(" / ")
        _add_field_run(p, 'NUMPAGES')


def build_docx(customer_name: str, website: str, monthly_budget: int, strategy_core: str, execution_text: str) -> bytes:
    """
    Simpel DOCX-builder, der tager den genererede strategi (strategy_core + execution_text)
    og skriver den ud med overskrifter og bullets i samme rækkefølge som i outputtet.
    """
    doc = Document()

    # Titel og basisinfo
    doc.add_heading("Meta Ads Strategi", 0)
    p = doc.add_paragraph()
    p.add_run(f"{customer_name or 'Kunde'} — {website or ''}").font.size = Pt(12)
    if monthly_budget:
        p = doc.add_paragraph()
        p.add_run(f"Månedligt budget: {monthly_budget} DKK").font.size = Pt(12)

    # Sidetal i footer
    add_page_numbers(doc)

    # Kendte sektionstitler i den nye struktur
    section_headers = [
        "Agenda",
        "Forretningen",
        "Introduktion til Meta strategi",
        "Målsætninger & KPI’er",
        "Målgruppe",
        "Kampagneplan & budgetplan",
        "USP’er & budskaber",
        "Content",
        "Next",
    ]
    campaign_df = build_campaign_table(execution_text or "")
    budget_df = build_budget_table(execution_text or "")

    full_text = ((strategy_core or "").strip() + "\n\n" + (execution_text or "").strip()).strip()

    if full_text:
        lines = full_text.splitlines()
        n = len(lines)

        # Identificér linjer der tilhører den måned-for-måned budgetblok,
        # så vi kan springe dem over i DOCX (budget vises kun som tabel).
        skip_indices = set()
        if not budget_df.empty:
            months = {
                "Januar",
                "Februar",
                "Marts",
                "April",
                "Maj",
                "Juni",
                "Juli",
                "August",
                "September",
                "Oktober",
                "November",
                "December",
            }
            budget_idx = None
            for i, ln in enumerate(lines):
                if ln.strip().lower().startswith("budgetplan"):
                    budget_idx = i
                    # spring også selve "Budgetplan:"-linjen over
                    skip_indices.add(i)
                    break

            if budget_idx is not None:
                i = budget_idx + 1
                while i < n:
                    s = lines[i].strip()
                    if not s:
                        skip_indices.add(i)
                        i += 1
                        continue
                    low = s.lower()
                    # Måned, "Budget (DKK):", "Fokus:", "Begrundelse:" skal ikke med som tekst
                    if s in months or low.startswith("budget") or low.startswith("fokus") or low.startswith("begrundelse"):
                        skip_indices.add(i)
                        i += 1
                        continue
                    # Stop når vi rammer fx "Opsummering:" eller næste sektion
                    break

        for idx, raw_line in enumerate(lines):
            if idx in skip_indices:
                continue

            line = raw_line.rstrip()
            if not line.strip():
                doc.add_paragraph()
                continue

            stripped = line.strip()

            # Hovedoverskrifter
            if stripped in section_headers:
                doc.add_heading(stripped, level=1)
                if stripped == "Kampagneplan & budgetplan" and not campaign_df.empty:
                    table = doc.add_table(rows=1 + len(campaign_df), cols=5)
                    table.style = "Table Grid"
                    hdr_cells = table.rows[0].cells
                    headers = ["Kampagne", "Formål", "Målgruppe", "Periode", "Budget"]
                    for i, h in enumerate(headers):
                        run = hdr_cells[i].paragraphs[0].add_run(h)
                        run.bold = True
                    for row_idx, row in enumerate(campaign_df.itertuples(index=False), start=1):
                        table.rows[row_idx].cells[0].text = getattr(row, "Kampagne", "")
                        table.rows[row_idx].cells[1].text = getattr(row, "Formål", "")
                        table.rows[row_idx].cells[2].text = getattr(row, "Målgruppe", "")
                        table.rows[row_idx].cells[3].text = getattr(row, "Periode", "")
                        table.rows[row_idx].cells[4].text = getattr(row, "Budget", "")
                    doc.add_paragraph()
                    if not budget_df.empty:
                        btable = doc.add_table(rows=1 + len(budget_df), cols=4)
                        btable.style = "Table Grid"
                        bhdr = btable.rows[0].cells
                        bheaders = ["Måned", "Budget (DKK)", "Fokus", "Begrundelse"]
                        for i, h in enumerate(bheaders):
                            run = bhdr[i].paragraphs[0].add_run(h)
                            run.bold = True
                        for row_idx, row in enumerate(budget_df.itertuples(index=False), start=1):
                            btable.rows[row_idx].cells[0].text = getattr(row, "Måned", "")
                            btable.rows[row_idx].cells[1].text = getattr(row, "Budget (DKK)", "")
                            btable.rows[row_idx].cells[2].text = getattr(row, "Fokus", "")
                            btable.rows[row_idx].cells[3].text = getattr(row, "Begrundelse", "")
                        doc.add_paragraph()
            # Bullets
            elif stripped.startswith("- ") or stripped.startswith("• "):
                doc.add_paragraph(stripped[2:].strip(), style="List Bullet")
            else:
                doc.add_paragraph(stripped)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


#
# =============================
# Hoved-flow (Streamlit)
# =============================

if generate_btn:
    if not api_key:
        st.error("OPENAI_API_KEY mangler i .streamlit/secrets.toml")
        st.stop()

    ctx = build_context()
    strategy_core = ""
    execution_text = ""

    # Kald 1: Strategi-kernen (streaming)
    with st.spinner("Strategi-kernen…"):
        strategy_placeholder = st.empty()
        strategy_core_raw = run_gpt(
            prompt_strategy_core(ctx),
            api_key,
            model,
            max_tokens=3200,
            stream_placeholder=strategy_placeholder,
        )
        if strategy_core_raw.startswith("[AI-ERROR]") or strategy_core_raw.startswith("[FEJL]"):
            st.error(strategy_core_raw)
        strategy_placeholder.empty()
        st.subheader("Strategi")
        strategy_core = sanitize(strategy_core_raw)
        formatted_core = format_headings(strategy_core)
        st.markdown(formatted_core.replace("\n", "<br>"), unsafe_allow_html=True)

    # Kald 2: Eksekvering (baseret på kernen, streaming)
    with st.spinner("Eksekvering…"):
        execution_placeholder = st.empty()
        execution_raw = run_gpt(
            prompt_execution(ctx, strategy_core),
            api_key,
            model,
            max_tokens=3200,
            stream_placeholder=execution_placeholder,
        )
        if execution_raw.startswith("[AI-ERROR]") or execution_raw.startswith("[FEJL]"):
            st.error(execution_raw)
        execution_placeholder.empty()
        st.subheader("Eksekvering")
        execution_text = sanitize(execution_raw)

        # Byg tabeller
        campaign_df = build_campaign_table(execution_text)
        budget_df = build_budget_table(execution_text)

        lines = execution_text.splitlines()
        n = len(lines)
        idx_kamp = None
        idx_budget = None

        months = {
            "Januar",
            "Februar",
            "Marts",
            "April",
            "Maj",
            "Juni",
            "Juli",
            "August",
            "September",
            "Oktober",
            "November",
            "December",
        }

        # Find "Kampagneoversigt" og "Budgetplan"
        for i, ln in enumerate(lines):
            if idx_kamp is None and ln.strip() == "Kampagneoversigt":
                idx_kamp = i
            if idx_budget is None and ln.strip().lower().startswith("budgetplan"):
                idx_budget = i
            if idx_kamp is not None and idx_budget is not None:
                break

        pointer = 0

        # Tekst op til (og inkl.) "Kampagneoversigt" + kampagnetabel
        if idx_kamp is not None:
            top_block = "\n".join(lines[: idx_kamp + 1])
            if top_block:
                formatted_top = format_headings(top_block)
                st.markdown(formatted_top.replace("\n", "<br>"), unsafe_allow_html=True)

            if not campaign_df.empty:
                st.subheader("📊 Kampagneoversigt")
                st.dataframe(campaign_df, use_container_width=True)

            # VIGTIGT: vi springer kampagnebeskrivelserne helt over
            pointer = idx_budget if idx_budget is not None else idx_kamp + 1
        else:
            pointer = 0

        # Budgetplan + tabel (ingen måned-for-måned tekst)
        if idx_budget is not None and idx_budget < n:
            # Vis selve "Budgetplan:"-linjen
            st.markdown(lines[idx_budget].replace("\n", "  \n"), unsafe_allow_html=True)
            pointer = idx_budget + 1

            if not budget_df.empty:
                st.subheader("📈 Budgetplan")
                st.dataframe(budget_df, use_container_width=True)

            # Spring måneds-blokkene over i tekstvisningen
            summary_idx = None
            for i in range(pointer, n):
                s = lines[i].strip()
                if s.lower().startswith("opsummering"):
                    summary_idx = i
                    break

            if summary_idx is not None:
                pointer = summary_idx
            else:
                # Hvis der ikke er en 'Opsummering'-linje, så skip alle linjer der ligner måned/budget/fokus/begrundelse
                i = pointer
                while i < n:
                    s = lines[i].strip()
                    if not s:
                        i += 1
                        continue
                    low = s.lower()
                    if s in months or low.startswith("budget") or low.startswith("fokus") or low.startswith("begrundelse"):
                        i += 1
                        continue
                    break
                pointer = i

        # Resten af teksten (efter budget-blokken eller – hvis ingen budget – efter kampagnedelen)
        if pointer < n:
            rest_block = "\n".join(lines[pointer:])
            if rest_block:
                formatted_rest = format_headings(rest_block)
                st.markdown(formatted_rest.replace("\n", "<br>"), unsafe_allow_html=True)

    # DOCX-export
    try:
        docx_bytes = build_docx(customer_name, website, monthly_budget, strategy_core, execution_text)
        st.success("Word-dokument klar – hent herunder")
        st.download_button(
            "⬇️ Download DOCX",
            data=docx_bytes,
            file_name=f"meta_strategi_{(customer_name or 'kunde').replace(' ', '_')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    except Exception as e:
        st.error(f"DOCX-export fejlede: {e}")